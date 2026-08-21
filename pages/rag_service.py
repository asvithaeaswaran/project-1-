"""
RAG (Retrieval-Augmented Generation) Service Module
- Embeddings: sentence-transformers/all-MiniLM-L6-v2
- Vector Store: FAISS (Facebook AI Similarity Search)
- Framework: LangChain
- LLMs: Google Gemini, Ollama (Local), Groq, OpenAI, and Built-in Extractive Fallback
"""

import os
import re
import csv
import json
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple

from django.conf import settings
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)

# Global singleton cache for embedding model to prevent reload overhead
_EMBEDDINGS_INSTANCE = None


def get_embeddings_model():
    """
    Get or initialize the singleton HuggingFace all-MiniLM-L6-v2 embeddings model.
    """
    global _EMBEDDINGS_INSTANCE
    if _EMBEDDINGS_INSTANCE is None:
        try:
            from langchain_huggingface import HuggingFaceEmbeddings
            _EMBEDDINGS_INSTANCE = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={"device": "cpu"},
                encode_kwargs={"normalize_embeddings": True}
            )
            logger.info("Successfully loaded all-MiniLM-L6-v2 embeddings.")
        except Exception as e:
            logger.error(f"Error loading HuggingFace embeddings: {e}")
            from langchain_community.embeddings import HuggingFaceEmbeddings
            _EMBEDDINGS_INSTANCE = HuggingFaceEmbeddings(
                model_name="all-MiniLM-L6-v2"
            )
    return _EMBEDDINGS_INSTANCE


def get_user_vectorstore_dir(user_id) -> Path:
    """
    Returns the FAISS vector store directory path for a specific user.
    """
    base_dir = Path(settings.MEDIA_ROOT) / 'vectorstores' / f'user_{user_id}'
    base_dir.mkdir(parents=True, exist_ok=True)
    return base_dir


# ==============================================================================
# DOCUMENT PARSING & TEXT EXTRACTION
# ==============================================================================

def extract_text_from_pdf(file_path: str) -> List[Tuple[str, int]]:
    """
    Extract text from PDF file page by page.
    Returns list of (page_text, page_number) tuples.
    """
    results = []
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                results.append((text, i + 1))
    except Exception as e:
        logger.error(f"Error reading PDF {file_path}: {e}")
        try:
            import pypdf
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for i, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        results.append((text.strip(), i + 1))
        except Exception:
            pass
    return results


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file.
    """
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
                    
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Error reading DOCX {file_path}: {e}")
        return ""


def extract_text_from_csv(file_path: str) -> str:
    """
    Extract structured text from CSV file.
    """
    lines = []
    try:
        with open(file_path, mode='r', encoding='utf-8', errors='replace') as f:
            reader = csv.reader(f)
            header = None
            for row in reader:
                if not row or not any(row):
                    continue
                if header is None:
                    header = row
                    lines.append(f"Columns: {', '.join(header)}")
                else:
                    item_str = ", ".join(f"{h}: {v}" for h, v in zip(header, row) if v.strip())
                    if item_str:
                        lines.append(item_str)
    except Exception as e:
        logger.error(f"Error reading CSV {file_path}: {e}")
    return "\n".join(lines)


def extract_text_from_plain(file_path: str) -> str:
    """
    Extract text from plain text or markdown files with auto encoding detection.
    """
    for enc in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.error(f"Error reading plain file {file_path}: {e}")
            break
    return ""


def parse_and_chunk_document(doc_model) -> List[Document]:
    """
    Parses an UploadedDocument instance and splits it into LangChain Document chunks.
    """
    file_path = doc_model.file.path
    file_type = doc_model.file_type.lower()
    doc_id = str(doc_model.id)
    doc_title = doc_model.title
    filename = doc_model.original_filename

    raw_docs: List[Document] = []

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=650,
        chunk_overlap=120,
        separators=["\n\n", "\n", ". ", "; ", ", ", " ", ""],
        length_function=len
    )

    if file_type == 'pdf':
        page_tuples = extract_text_from_pdf(file_path)
        if not page_tuples:
            full_text = extract_text_from_plain(file_path)
            if full_text:
                page_tuples = [(full_text, 1)]

        for page_text, page_num in page_tuples:
            chunks = text_splitter.split_text(page_text)
            for idx, chunk in enumerate(chunks):
                raw_docs.append(Document(
                    page_content=chunk,
                    metadata={
                        "doc_id": doc_id,
                        "doc_title": doc_title,
                        "filename": filename,
                        "page": page_num,
                        "chunk_index": idx,
                        "file_type": file_type
                    }
                ))

    elif file_type in ['docx', 'doc']:
        full_text = extract_text_from_docx(file_path)
        chunks = text_splitter.split_text(full_text)
        for idx, chunk in enumerate(chunks):
            raw_docs.append(Document(
                page_content=chunk,
                metadata={
                    "doc_id": doc_id,
                    "doc_title": doc_title,
                    "filename": filename,
                    "page": 1,
                    "chunk_index": idx,
                    "file_type": file_type
                }
            ))

    elif file_type == 'csv':
        full_text = extract_text_from_csv(file_path)
        chunks = text_splitter.split_text(full_text)
        for idx, chunk in enumerate(chunks):
            raw_docs.append(Document(
                page_content=chunk,
                metadata={
                    "doc_id": doc_id,
                    "doc_title": doc_title,
                    "filename": filename,
                    "page": 1,
                    "chunk_index": idx,
                    "file_type": file_type
                }
            ))

    else:
        full_text = extract_text_from_plain(file_path)
        chunks = text_splitter.split_text(full_text)
        for idx, chunk in enumerate(chunks):
            raw_docs.append(Document(
                page_content=chunk,
                metadata={
                    "doc_id": doc_id,
                    "doc_title": doc_title,
                    "filename": filename,
                    "page": 1,
                    "chunk_index": idx,
                    "file_type": file_type
                }
            ))

    return raw_docs


# ==============================================================================
# FAISS VECTOR STORE OPERATIONS
# ==============================================================================

def index_uploaded_document(doc_model) -> Tuple[bool, int, Optional[str]]:
    """
    Parses, chunks, and indexes a single UploadedDocument into the user's FAISS index.
    Returns: (success, chunk_count, error_message)
    """
    try:
        from langchain_community.vectorstores import FAISS

        chunks = parse_and_chunk_document(doc_model)
        if not chunks:
            return False, 0, "No readable text could be extracted from the document."

        embeddings = get_embeddings_model()
        user_id = doc_model.user.id if doc_model.user else 0
        vectorstore_dir = get_user_vectorstore_dir(user_id)
        index_file = vectorstore_dir / "index.faiss"

        if index_file.exists():
            try:
                vectorstore = FAISS.load_local(
                    str(vectorstore_dir),
                    embeddings,
                    allow_dangerous_deserialization=True
                )
                vectorstore.add_documents(chunks)
            except Exception as load_err:
                logger.warning(f"Error loading existing index, recreating: {load_err}")
                vectorstore = FAISS.from_documents(chunks, embeddings)
        else:
            vectorstore = FAISS.from_documents(chunks, embeddings)

        vectorstore.save_local(str(vectorstore_dir))

        total_chars = sum(len(c.page_content) for c in chunks)
        doc_model.chunk_count = len(chunks)
        doc_model.char_count = total_chars
        doc_model.status = 'ready'
        doc_model.error_message = None
        doc_model.save()

        return True, len(chunks), None

    except Exception as e:
        logger.exception(f"Error indexing document {doc_model.id}: {e}")
        doc_model.status = 'error'
        doc_model.error_message = str(e)
        doc_model.save()
        return False, 0, str(e)


def rebuild_user_vectorstore(user) -> bool:
    """
    Rebuilds the complete FAISS index for a user from all their active 'ready' documents.
    Used when a document is deleted.
    """
    try:
        from langchain_community.vectorstores import FAISS
        from .models import UploadedDocument

        user_id = user.id if user else 0
        vectorstore_dir = get_user_vectorstore_dir(user_id)

        for item in vectorstore_dir.glob("*"):
            try:
                if item.is_file():
                    item.unlink()
            except Exception:
                pass

        documents = UploadedDocument.objects.filter(user=user, status='ready')
        all_chunks: List[Document] = []

        for doc in documents:
            if os.path.exists(doc.file.path):
                chunks = parse_and_chunk_document(doc)
                all_chunks.extend(chunks)

        if all_chunks:
            embeddings = get_embeddings_model()
            vectorstore = FAISS.from_documents(all_chunks, embeddings)
            vectorstore.save_local(str(vectorstore_dir))

        return True
    except Exception as e:
        logger.error(f"Error rebuilding vector store for user {user}: {e}")
        return False


def load_user_vectorstore(user_id):
    """
    Loads FAISS vector store for the user if exists, else returns None.
    """
    from langchain_community.vectorstores import FAISS
    vectorstore_dir = get_user_vectorstore_dir(user_id)
    index_file = vectorstore_dir / "index.faiss"

    if not index_file.exists():
        return None

    try:
        embeddings = get_embeddings_model()
        return FAISS.load_local(
            str(vectorstore_dir),
            embeddings,
            allow_dangerous_deserialization=True
        )
    except Exception as e:
        logger.error(f"Failed to load FAISS index for user {user_id}: {e}")
        return None


def search_documents(query: str, user, doc_ids: Optional[List[str]] = None, k: int = 4) -> List[Tuple[Document, float]]:
    """
    Searches the user's FAISS index for the most relevant document chunks.
    Filters by doc_ids if specified.
    """
    user_id = user.id if user else 0
    vectorstore = load_user_vectorstore(user_id)
    if not vectorstore:
        return []

    try:
        fetch_k = k * 3 if doc_ids else k
        results_with_scores = vectorstore.similarity_search_with_score(query, k=fetch_k)

        filtered_results = []
        for doc, score in results_with_scores:
            if doc_ids:
                if str(doc.metadata.get("doc_id")) in doc_ids:
                    filtered_results.append((doc, float(score)))
            else:
                filtered_results.append((doc, float(score)))

            if len(filtered_results) >= k:
                break

        return filtered_results
    except Exception as e:
        logger.error(f"Error performing similarity search: {e}")
        return []


# ==============================================================================
# LLM INSTANTIATION & RAG GENERATION
# ==============================================================================

def get_rag_llm(request=None, custom_provider=None, custom_model=None):
    """
    Initializes and returns a configured LangChain Chat Model based on user session/env.
    Supported: Google Gemini, Ollama, Groq, OpenAI.
    """
    api_key = None
    if request:
        api_key = request.session.get('gemini_api_key') or request.session.get('openai_api_key') or request.session.get('api_key')
        if not custom_provider:
            custom_provider = request.session.get('ai_provider')
        if not custom_model:
            custom_model = request.session.get('gemini_model') or request.session.get('openai_model') or request.session.get('selected_model')

    if not api_key:
        api_key = os.environ.get('GEMINI_API_KEY') or os.environ.get('GOOGLE_API_KEY') or os.environ.get('OPENAI_API_KEY') or os.environ.get('GROQ_API_KEY')

    ollama_url = (request.session.get('ollama_base_url') if request else None) or os.environ.get('OLLAMA_BASE_URL', 'http://localhost:11434')

    # 1. Check if Ollama requested or auto-configured
    if custom_provider == 'ollama' or (custom_model and ('llama' in custom_model or 'mistral' in custom_model or 'gemma' in custom_model or 'phi' in custom_model or 'deepseek' in custom_model) and not api_key):
        try:
            from langchain_ollama import ChatOllama
            model_name = custom_model or 'llama3'
            return ChatOllama(
                model=model_name,
                base_url=ollama_url,
                temperature=0.3
            ), f"Ollama Local ({model_name})"
        except Exception as e:
            logger.warning(f"Failed to initialize ChatOllama: {e}")

    # 2. Check if Google Gemini requested or API key starts with AIza
    if (custom_provider == 'gemini') or (api_key and api_key.startswith('AIza')) or (custom_model and 'gemini' in str(custom_model).lower()):
        try:
            from langchain_google_genai import ChatGoogleGenerativeAI
            model_name = custom_model if custom_model and 'gemini' in custom_model else 'gemini-2.0-flash'
            gemini_key = api_key if api_key and api_key.startswith('AIza') else os.environ.get('GEMINI_API_KEY') or os.environ.get('GOOGLE_API_KEY') or api_key
            return ChatGoogleGenerativeAI(
                model=model_name,
                google_api_key=gemini_key,
                temperature=0.3,
                max_output_tokens=2048
            ), f"Google Gemini ({model_name})"
        except Exception as e:
            logger.warning(f"Failed to initialize ChatGoogleGenerativeAI: {e}")

    # 3. Groq (starts with gsk_ or configured)
    if (api_key and api_key.startswith('gsk_')) or custom_provider == 'groq':
        try:
            from langchain_openai import ChatOpenAI
            model_name = custom_model if custom_model and '/' in custom_model else 'openai/gpt-oss-120b'
            return ChatOpenAI(
                model=model_name,
                api_key=api_key or os.environ.get('OPENAI_API_KEY'),
                base_url="https://api.groq.com/openai/v1",
                temperature=0.3
            ), f"Groq AI ({model_name})"
        except Exception as e:
            logger.warning(f"Failed to initialize Groq Chat: {e}")

    # 4. Standard OpenAI or OpenRouter
    if api_key and (api_key.startswith('sk-') or api_key.startswith('sk-proj-') or api_key.startswith('sk-or-')):
        try:
            from langchain_openai import ChatOpenAI
            base_url = "https://openrouter.ai/api/v1" if api_key.startswith('sk-or-') else "https://api.openai.com/v1"
            model_name = custom_model or ('meta-llama/llama-3.3-70b-instruct:free' if api_key.startswith('sk-or-') else 'gpt-4o-mini')
            return ChatOpenAI(
                model=model_name,
                api_key=api_key,
                base_url=base_url,
                temperature=0.3
            ), f"OpenAI ({model_name})"
        except Exception as e:
            logger.warning(f"Failed to initialize ChatOpenAI: {e}")

    # 5. Default attempt with any available key in environment
    if os.environ.get('OPENAI_API_KEY'):
        try:
            from langchain_openai import ChatOpenAI
            base_url = os.environ.get('OPENAI_BASE_URL', 'https://api.groq.com/openai/v1')
            model_name = os.environ.get('OPENAI_MODEL', 'openai/gpt-oss-120b')
            return ChatOpenAI(
                model=model_name,
                api_key=os.environ.get('OPENAI_API_KEY'),
                base_url=base_url,
                temperature=0.3
            ), f"Live AI ({model_name})"
        except Exception:
            pass

    return None, "Built-in Extractive Engine"


def generate_rag_answer(query: str, user, doc_ids: Optional[List[str]] = None, conversation_history: Optional[List[Dict[str, str]]] = None, request=None) -> Dict[str, Any]:
    """
    Complete LangChain RAG pipeline:
    1. Embeds query with all-MiniLM-L6-v2
    2. Retrieves top-k relevant chunks from FAISS vector store
    3. Formats prompt template with document context and citations
    4. Generates synthesized response using Google Gemini, Ollama, Groq, or built-in engine
    5. Returns answer with detailed source citations
    """
    relevant_chunks_with_scores = search_documents(query, user, doc_ids=doc_ids, k=4)

    if not relevant_chunks_with_scores:
        return {
            "answer": "I could not find any relevant information in your uploaded documents regarding your query. Please make sure documents are uploaded and processed in your Knowledge Base.",
            "sources": [],
            "provider": "FAISS Search",
            "chunks_count": 0,
            "status": "no_context"
        }

    sources_data = []
    context_blocks = []

    for rank, (chunk, score) in enumerate(relevant_chunks_with_scores, 1):
        doc_id = chunk.metadata.get("doc_id")
        doc_title = chunk.metadata.get("doc_title", "Document")
        page = chunk.metadata.get("page", 1)
        filename = chunk.metadata.get("filename", "")
        content = chunk.page_content.strip()

        sources_data.append({
            "rank": rank,
            "doc_id": doc_id,
            "doc_title": doc_title,
            "filename": filename,
            "page": page,
            "score": round(score, 4),
            "snippet": content[:300] + ("..." if len(content) > 300 else "")
        })

        context_blocks.append(
            f"--- EXCERPT {rank} [Source: {doc_title}, Page/Section: {page}] ---\n{content}"
        )

    full_context = "\n\n".join(context_blocks)

    system_prompt = (
        "You are an expert AI Document Analysis and Research Assistant. "
        "Your task is to answer the user's question accurately and thoroughly based ONLY on the provided document excerpts below.\n\n"
        "Guidelines:\n"
        "1. Base your answer strictly on the provided context.\n"
        "2. If the context does not contain enough information to answer fully, explain what is available and clarify what is missing.\n"
        "3. Provide clean, well-formatted Markdown with bullet points, bold key terms, and numbered steps where appropriate.\n"
        "4. Include direct references to the document titles/sections when quoting facts or numbers."
    )

    user_rag_prompt = (
        f"DOCUMENT EXCERPTS:\n{full_context}\n\n"
        f"USER QUESTION: {query}\n\n"
        "Please provide a comprehensive, accurate answer based on the document excerpts above."
    )

    llm, provider_name = get_rag_llm(request)

    if llm is not None:
        try:
            from langchain_core.messages import SystemMessage, HumanMessage

            messages = [
                SystemMessage(content=system_prompt),
            ]

            if conversation_history:
                for h in conversation_history[-4:]:
                    if h.get('role') == 'user':
                        messages.append(HumanMessage(content=h.get('content', '')))
                    elif h.get('role') == 'assistant':
                        from langchain_core.messages import AIMessage
                        messages.append(AIMessage(content=h.get('content', '')))

            messages.append(HumanMessage(content=user_rag_prompt))

            response = llm.invoke(messages)
            answer_content = response.content if hasattr(response, 'content') else str(response)

            return {
                "answer": answer_content,
                "sources": sources_data,
                "provider": provider_name,
                "chunks_count": len(sources_data),
                "status": "success"
            }

        except Exception as e:
            logger.error(f"Error during LLM invocation ({provider_name}): {e}")
            fallback_answer = generate_extractive_fallback_answer(query, relevant_chunks_with_scores)
            return {
                "answer": fallback_answer + f"\n\n> ℹ️ *Note: Answer synthesized using local FAISS vector search due to LLM connection warning: `{str(e)[:120]}`*",
                "sources": sources_data,
                "provider": f"{provider_name} (Extractive)",
                "chunks_count": len(sources_data),
                "status": "fallback"
            }

    fallback_answer = generate_extractive_fallback_answer(query, relevant_chunks_with_scores)
    return {
        "answer": fallback_answer,
        "sources": sources_data,
        "provider": "FAISS RAG Engine",
        "chunks_count": len(sources_data),
        "status": "success"
    }


def generate_extractive_fallback_answer(query: str, chunks_with_scores: List[Tuple[Document, float]]) -> str:
    """
    Extractive answer generator that finds and formats the most relevant statements from the retrieved chunks.
    """
    top_chunks = [c[0] for c in chunks_with_scores]
    lines = [f"Based on your documents, here are the most relevant findings for **\"{query}\"**:\n"]

    for i, chunk in enumerate(top_chunks[:3], 1):
        title = chunk.metadata.get("doc_title", "Document")
        page = chunk.metadata.get("page", 1)
        text = chunk.page_content.strip()

        clean_text = " ".join(text.split())
        if len(clean_text) > 400:
            clean_text = clean_text[:397] + "..."

        lines.append(f"### {i}. Source: *{title}* (Page/Section {page})")
        lines.append(f"> \"{clean_text}\"\n")

    lines.append("---")
    lines.append("💡 *You can configure Google Gemini or Ollama local model in Settings for generative conversational synthesis.*")
    return "\n".join(lines)


# ==============================================================================
# OLLAMA LOCAL MODEL DISCOVERY HELPER
# ==============================================================================

def check_ollama_status(base_url="http://localhost:11434") -> Dict[str, Any]:
    """
    Pings local Ollama service to check if it is running and retrieve list of available local models.
    """
    import urllib.request
    clean_url = base_url.rstrip('/')
    tags_url = f"{clean_url}/api/tags"

    try:
        req = urllib.request.Request(tags_url, headers={'User-Agent': 'Gelato-RAG/1.0'})
        with urllib.request.urlopen(req, timeout=2.5) as response:
            if response.status == 200:
                data = json.loads(response.read().decode('utf-8'))
                models = [m.get('name') for m in data.get('models', []) if m.get('name')]
                return {
                    "running": True,
                    "base_url": clean_url,
                    "models": models,
                    "message": f"Ollama is online with {len(models)} local models."
                }
    except Exception as e:
        return {
            "running": False,
            "base_url": clean_url,
            "models": [],
            "message": f"Ollama not reachable at {clean_url} ({str(e)})"
        }
