import os
import csv
import json
import logging
from io import BytesIO

logger = logging.getLogger(__name__)

def extract_text_from_file(file_obj_or_path, filename=None):
    """
    Directly extracts clean text from PDF, DOCX, TXT, CSV, JSON, Markdown, and Code files.
    Works synchronously and instantly without external vector databases or heavy models.
    """
    if filename is None:
        if isinstance(file_obj_or_path, str):
            filename = os.path.basename(file_obj_or_path)
        elif hasattr(file_obj_or_path, 'name'):
            filename = file_obj_or_path.name
        else:
            filename = 'document.txt'

    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else 'txt'
    extracted_text = ""

    try:
        # Read content bytes if file object or path
        if isinstance(file_obj_or_path, str):
            with open(file_obj_or_path, 'rb') as f:
                file_bytes = f.read()
        elif hasattr(file_obj_or_path, 'read'):
            file_bytes = file_obj_or_path.read()
            if hasattr(file_obj_or_path, 'seek'):
                file_obj_or_path.seek(0)
        else:
            file_bytes = bytes(file_obj_or_path)

        # 1. PDF Extraction
        if ext == 'pdf':
            try:
                import pypdf
                reader = pypdf.PdfReader(BytesIO(file_bytes))
                pages_text = []
                for i, page in enumerate(reader.pages):
                    t = page.extract_text()
                    if t and t.strip():
                        pages_text.append(f"--- Page {i + 1} ---\n{t.strip()}")
                extracted_text = "\n\n".join(pages_text)
            except Exception as e:
                logger.error(f"Error extracting PDF text with pypdf: {e}")
                extracted_text = f"[PDF Extraction Error: {str(e)}]"

        # 2. DOCX Word Document Extraction
        elif ext in ['docx', 'doc']:
            try:
                import docx
                doc = docx.Document(BytesIO(file_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # Also extract table text if present
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            paragraphs.append(row_text)
                extracted_text = "\n\n".join(paragraphs)
            except Exception as e:
                logger.error(f"Error extracting DOCX text: {e}")
                extracted_text = f"[DOCX Extraction Error: {str(e)}]"

        # 3. CSV File Extraction
        elif ext == 'csv':
            try:
                text_content = file_bytes.decode('utf-8', errors='replace')
                reader = csv.reader(text_content.splitlines())
                rows = []
                for row in reader:
                    if row:
                        rows.append(" | ".join(row))
                extracted_text = "\n".join(rows)
            except Exception as e:
                extracted_text = file_bytes.decode('utf-8', errors='replace')

        # 4. JSON File Extraction
        elif ext == 'json':
            try:
                data = json.loads(file_bytes.decode('utf-8', errors='replace'))
                extracted_text = json.dumps(data, indent=2)
            except Exception:
                extracted_text = file_bytes.decode('utf-8', errors='replace')

        # 5. TXT, MD, Python, JS, HTML, etc.
        else:
            try:
                extracted_text = file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    extracted_text = file_bytes.decode('utf-8-sig')
                except UnicodeDecodeError:
                    extracted_text = file_bytes.decode('latin-1', errors='replace')

    except Exception as e:
        logger.error(f"Failed to extract document {filename}: {e}")
        extracted_text = f"[Failed to read document content: {str(e)}]"

    cleaned_text = extracted_text.strip()
    return {
        'filename': filename,
        'file_type': ext,
        'char_count': len(cleaned_text),
        'text': cleaned_text,
        'preview': cleaned_text[:300] + ('...' if len(cleaned_text) > 300 else '')
    }
