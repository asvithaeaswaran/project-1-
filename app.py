"""
==============================================================================
Simple RAG Document Q&A Application (Flask + LangChain + FAISS + all-MiniLM-L6-v2)
==============================================================================
A clean, self-contained RAG application:
1. Upload documents (PDF, DOCX, TXT)
2. Chunk & Embed with sentence-transformers/all-MiniLM-L6-v2
3. Store in FAISS vector database
4. Query using Google Gemini, local Ollama, Groq, or OpenAI with source citations
==============================================================================
"""

import os
import io
import shutil
from pathlib import Path
from flask import Flask, request, jsonify, render_template_string
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# LangChain & RAG imports
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

# App setup
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'simple_rag_uploads')
app.config['FAISS_FOLDER'] = os.path.join(os.path.dirname(__file__), 'simple_rag_faiss_index')
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['FAISS_FOLDER'], exist_ok=True)

# Global in-memory embedding model singleton
print("Initializing all-MiniLM-L6-v2 embeddings model...")
embeddings_model = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2",
    model_kwargs={"device": "cpu"},
    encode_kwargs={"normalize_embeddings": True}
)

# In-memory document tracker
uploaded_docs_registry = []


# ==============================================================================
# DOCUMENT PARSERS & CHUNKER
# ==============================================================================

def extract_text_from_file(file_path: str, filename: str) -> list[Document]:
    """Extract and chunk text from PDF, DOCX, or TXT."""
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else 'txt'
    raw_text = ""
    pages_data = []

    if ext == 'pdf':
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                txt = (page.extract_text() or "").strip()
                if txt:
                    pages_data.append((txt, i + 1))
        except Exception as e:
            print(f"Error extracting PDF: {e}")
    elif ext in ['docx', 'doc']:
        try:
            import docx
            doc = docx.Document(file_path)
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for t in doc.tables:
                for row in t.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paras.append(" | ".join(cells))
            raw_text = "\n\n".join(paras)
            pages_data = [(raw_text, 1)]
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
    else:
        # TXT / MD / CSV
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                raw_text = f.read()
            pages_data = [(raw_text, 1)]
        except Exception as e:
            print(f"Error extracting Text file: {e}")

    # Chunking
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=600,
        chunk_overlap=100,
        separators=["\n\n", "\n", ". ", " ", ""]
    )

    chunks = []
    for text_content, page_num in pages_data:
        split_texts = splitter.split_text(text_content)
        for idx, chunk in enumerate(split_texts):
            chunks.append(Document(
                page_content=chunk,
                metadata={
                    "filename": filename,
                    "page": page_num,
                    "chunk_index": idx
                }
            ))

    return chunks


def get_vectorstore():
    """Load or initialize FAISS vectorstore."""
    index_file = os.path.join(app.config['FAISS_FOLDER'], 'index.faiss')
    if os.path.exists(index_file):
        try:
            return FAISS.load_local(
                app.config['FAISS_FOLDER'],
                embeddings_model,
                allow_dangerous_deserialization=True
            )
        except Exception as e:
            print(f"Could not load existing index: {e}")
    return None


# ==============================================================================
# LLM INVOCATION HELPER (Gemini, Ollama, Groq, OpenAI)
# ==============================================================================

def invoke_llm(prompt: str, system_prompt: str, provider: str = "auto") -> str:
    """Call Google Gemini, Ollama, Groq, or OpenAI."""
    gemini_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    openai_key = os.environ.get("OPENAI_API_KEY")
    ollama_url = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434")

    # 1. Google Gemini
    if (provider == "gemini" or (not provider and gemini_key)) and gemini_key:
        try:
            from langchain_google_genai import ChatGoogleGenerativeAI
            from langchain_core.messages import SystemMessage, HumanMessage
            llm = ChatGoogleGenerativeAI(
                model=os.environ.get("GEMINI_MODEL", "gemini-2.0-flash"),
                google_api_key=gemini_key,
                temperature=0.3
            )
            resp = llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=prompt)])
            return resp.content
        except Exception as e:
            print(f"Gemini error: {e}")

    # 2. Ollama Local LLM
    if provider == "ollama":
        try:
            from langchain_ollama import ChatOllama
            from langchain_core.messages import SystemMessage, HumanMessage
            llm = ChatOllama(
                model=os.environ.get("OLLAMA_MODEL", "llama3"),
                base_url=ollama_url,
                temperature=0.3
            )
            resp = llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=prompt)])
            return resp.content
        except Exception as e:
            print(f"Ollama error: {e}")

    # 3. Groq or OpenAI
    if openai_key:
        try:
            from langchain_openai import ChatOpenAI
            from langchain_core.messages import SystemMessage, HumanMessage
            base_url = os.environ.get("OPENAI_BASE_URL", "https://api.groq.com/openai/v1" if openai_key.startswith("gsk_") else "https://api.openai.com/v1")
            model = os.environ.get("OPENAI_MODEL", "openai/gpt-oss-120b" if openai_key.startswith("gsk_") else "gpt-4o-mini")
            llm = ChatOpenAI(
                model=model,
                api_key=openai_key,
                base_url=base_url,
                temperature=0.3
            )
            resp = llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=prompt)])
            return resp.content
        except Exception as e:
            print(f"OpenAI/Groq error: {e}")

    return None


# ==============================================================================
# FLASK WEB ROUTES & APIS
# ==============================================================================

HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>My RAG Document Assistant</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.5.1/css/all.min.css">
    <style>
        .markdown-body p { margin-bottom: 0.75rem; }
        .markdown-body ul { list-style-type: disc; padding-left: 1.25rem; margin-bottom: 0.75rem; }
        .markdown-body pre { background: #0f172a; color: #f8fafc; padding: 0.75rem; border-radius: 0.5rem; margin: 0.5rem 0; overflow-x: auto; }
        .markdown-body code { background: rgba(0,0,0,0.08); padding: 0.1rem 0.3rem; border-radius: 0.2rem; font-size: 0.85em; }
    </style>
</head>
<body class="bg-slate-900 text-slate-100 min-h-screen flex flex-col font-sans">
    <!-- Header -->
    <header class="bg-slate-800 border-b border-slate-700 px-6 py-4 flex items-center justify-between shadow-md">
        <div class="flex items-center gap-3">
            <div class="w-10 h-10 rounded-xl bg-gradient-to-tr from-indigo-500 to-purple-600 flex items-center justify-center text-white text-lg shadow">
                <i class="fa-solid fa-brain"></i>
            </div>
            <div>
                <h1 class="text-base font-bold text-white tracking-wide">My RAG Document Assistant</h1>
                <p class="text-xs text-slate-400">FAISS Vector Store &bull; all-MiniLM-L6-v2 Embeddings &bull; LangChain</p>
            </div>
        </div>
        <div class="flex items-center gap-2">
            <span class="px-3 py-1 rounded-full text-xs font-semibold bg-emerald-500/10 text-emerald-400 border border-emerald-500/20 flex items-center gap-1.5">
                <span class="w-2 h-2 rounded-full bg-emerald-400 animate-pulse"></span>
                RAG Engine Online
            </span>
        </div>
    </header>

    <div class="flex-1 max-w-7xl w-full mx-auto p-4 md:p-6 grid grid-cols-1 md:grid-cols-3 gap-6">
        <!-- Left Column: Document Upload & Knowledge Base -->
        <div class="bg-slate-800 rounded-2xl border border-slate-700 p-5 flex flex-col h-[78vh] shadow-xl">
            <h2 class="text-sm font-bold text-white mb-3 flex items-center gap-2">
                <i class="fa-solid fa-folder-open text-indigo-400"></i>
                <span>Knowledge Base</span>
            </h2>

            <!-- Upload Area -->
            <div id="dropzone" class="border-2 border-dashed border-slate-600 hover:border-indigo-500 bg-slate-900/50 rounded-xl p-5 text-center cursor-pointer transition mb-4">
                <input type="file" id="fileInput" multiple accept=".pdf,.docx,.txt" class="hidden">
                <i class="fa-solid fa-cloud-arrow-up text-3xl text-indigo-400 mb-2"></i>
                <div class="text-xs font-semibold text-white">Click or Drop PDF / DOCX / TXT</div>
                <div class="text-[10px] text-slate-400 mt-1">Embeds automatically into FAISS</div>
            </div>

            <!-- Upload status banner -->
            <div id="uploadStatus" class="hidden text-xs p-2.5 rounded-lg bg-indigo-950/60 border border-indigo-500/30 text-indigo-300 text-center mb-3">
                <i class="fa-solid fa-circle-notch fa-spin mr-1.5"></i> Indexing into FAISS...
            </div>

            <!-- Documents List -->
            <div class="flex items-center justify-between text-xs text-slate-400 mb-2">
                <span class="font-medium">Uploaded Files (<span id="docCount">0</span>)</span>
                <button id="clearBtn" class="text-rose-400 hover:text-rose-300 text-[11px] transition">Reset Index</button>
            </div>
            <div id="docsList" class="flex-1 overflow-y-auto space-y-2 pr-1 text-xs">
                <div class="text-center py-10 text-slate-500">No documents uploaded yet.</div>
            </div>
        </div>

        <!-- Right Column: RAG Chat Interface -->
        <div class="md:col-span-2 bg-slate-800 rounded-2xl border border-slate-700 p-5 flex flex-col h-[78vh] shadow-xl">
            <!-- Messages Container -->
            <div id="chatBox" class="flex-1 overflow-y-auto space-y-4 pr-2 mb-4">
                <div class="text-center py-16">
                    <div class="w-12 h-12 rounded-2xl bg-indigo-500/10 text-indigo-400 flex items-center justify-center text-2xl mx-auto mb-3">
                        <i class="fa-solid fa-comments"></i>
                    </div>
                    <h3 class="text-base font-bold text-white mb-1">Ask questions about your documents</h3>
                    <p class="text-xs text-slate-400 max-w-sm mx-auto">Upload documents on the left, then ask any question. Answers are synthesized using FAISS semantic search.</p>
                </div>
            </div>

            <!-- Typing Indicator -->
            <div id="typing" class="hidden text-xs text-indigo-400 mb-3 flex items-center gap-2">
                <i class="fa-solid fa-circle-notch fa-spin"></i>
                <span>Searching FAISS and synthesizing answer...</span>
            </div>

            <!-- Prompt Input Form -->
            <form id="queryForm" class="flex gap-2">
                <input
                    type="text"
                    id="queryInput"
                    placeholder="Ask a question about the uploaded documents..."
                    class="flex-1 bg-slate-900 border border-slate-700 rounded-xl px-4 py-3 text-xs text-white placeholder-slate-500 focus:outline-none focus:border-indigo-500"
                    required
                />
                <button
                    type="submit"
                    id="submitBtn"
                    class="px-5 py-3 rounded-xl bg-indigo-600 hover:bg-indigo-500 text-white font-semibold text-xs shadow-md transition flex items-center gap-2"
                >
                    <span>Ask</span>
                    <i class="fa-solid fa-arrow-up text-xs"></i>
                </button>
            </form>
        </div>
    </div>

    <script>
        const dropzone = document.getElementById('dropzone');
        const fileInput = document.getElementById('fileInput');
        const uploadStatus = document.getElementById('uploadStatus');
        const docsList = document.getElementById('docsList');
        const docCount = document.getElementById('docCount');
        const clearBtn = document.getElementById('clearBtn');
        const chatBox = document.getElementById('chatBox');
        const queryForm = document.getElementById('queryForm');
        const queryInput = document.getElementById('queryInput');
        const typing = document.getElementById('typing');
        const submitBtn = document.getElementById('submitBtn');

        let documents = [];

        dropzone.addEventListener('click', () => fileInput.click());
        fileInput.addEventListener('change', (e) => handleUpload(e.target.files));

        dropzone.addEventListener('dragover', (e) => { e.preventDefault(); dropzone.classList.add('border-indigo-500'); });
        dropzone.addEventListener('dragleave', () => dropzone.classList.remove('border-indigo-500'));
        dropzone.addEventListener('drop', (e) => {
            e.preventDefault();
            dropzone.classList.remove('border-indigo-500');
            handleUpload(e.dataTransfer.files);
        });

        async function handleUpload(files) {
            if (!files || files.length === 0) return;
            uploadStatus.classList.remove('hidden');

            const formData = new FormData();
            for (let i = 0; i < files.length; i++) {
                formData.append('files', files[i]);
            }

            try {
                const res = await fetch('/upload', { method: 'POST', body: formData });
                const data = await res.json();
                if (data.status === 'success') {
                    fetchDocs();
                } else {
                    alert('Upload failed: ' + data.error);
                }
            } catch (err) {
                alert('Upload error: ' + err.message);
            } finally {
                uploadStatus.classList.add('hidden');
                fileInput.value = '';
            }
        }

        async function fetchDocs() {
            try {
                const res = await fetch('/documents');
                const data = await res.json();
                documents = data.documents || [];
                docCount.textContent = documents.length;
                renderDocs();
            } catch (err) {
                console.error(err);
            }
        }

        function renderDocs() {
            docsList.innerHTML = '';
            if (documents.length === 0) {
                docsList.innerHTML = '<div class="text-center py-10 text-slate-500">No documents uploaded yet.</div>';
                return;
            }

            documents.forEach(doc => {
                const el = document.createElement('div');
                el.className = 'p-2.5 rounded-xl bg-slate-900/80 border border-slate-700 flex items-center justify-between text-xs';
                el.innerHTML = `
                    <div class="flex items-center gap-2 truncate">
                        <i class="fa-solid fa-file-lines text-indigo-400"></i>
                        <div class="truncate">
                            <div class="font-semibold text-white truncate">${doc.filename}</div>
                            <div class="text-[10px] text-slate-400">${doc.chunks} chunks embedded</div>
                        </div>
                    </div>
                    <span class="text-[10px] px-2 py-0.5 rounded-full bg-emerald-950/60 text-emerald-400 border border-emerald-500/20 font-medium">Ready</span>
                `;
                docsList.appendChild(el);
            });
        }

        clearBtn.addEventListener('click', async () => {
            if (confirm('Clear all uploaded documents and reset FAISS index?')) {
                await fetch('/clear', { method: 'POST' });
                fetchDocs();
                chatBox.innerHTML = '<div class="text-center py-16 text-xs text-slate-500">Index cleared.</div>';
            }
        });

        queryForm.addEventListener('submit', async (e) => {
            e.preventDefault();
            const text = queryInput.value.trim();
            if (!text) return;

            queryInput.value = '';
            appendBubble('user', text);
            typing.classList.remove('hidden');
            submitBtn.disabled = true;

            try {
                const res = await fetch('/query', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ query: text })
                });
                const data = await res.json();

                appendBubble('assistant', data.answer, data.sources);
            } catch (err) {
                appendBubble('assistant', '⚠️ Error: ' + err.message);
            } finally {
                typing.classList.add('hidden');
                submitBtn.disabled = false;
                chatBox.scrollTop = chatBox.scrollHeight;
            }
        });

        function appendBubble(role, content, sources = []) {
            const isUser = role === 'user';
            const wrapper = document.createElement('div');
            wrapper.className = `flex flex-col ${isUser ? 'items-end' : 'items-start'} text-xs space-y-1`;

            let sourcesHtml = '';
            if (!isUser && sources && sources.length > 0) {
                const items = sources.map((s, idx) => `
                    <div class="p-2 rounded bg-slate-900 border border-slate-700/80 text-[11px] mt-1">
                        <div class="font-semibold text-indigo-300 flex justify-between">
                            <span>Excerpt #${idx + 1} &bull; ${s.filename} (Page ${s.page})</span>
                            <span class="text-[10px] text-emerald-400">Score: ${s.score}</span>
                        </div>
                        <p class="text-slate-400 italic mt-1 line-clamp-2">"${s.snippet}"</p>
                    </div>
                `).join('');

                sourcesHtml = `
                    <details class="mt-2 pt-2 border-t border-slate-700/60 w-full text-slate-400">
                        <summary class="cursor-pointer font-medium text-indigo-400 hover:text-indigo-300 text-[11px]">
                            📚 View ${sources.length} Referenced Source Chunks
                        </summary>
                        <div class="mt-1.5 space-y-1.5">${items}</div>
                    </details>
                `;
            }

            wrapper.innerHTML = `
                <div class="text-[10px] font-semibold text-slate-400 px-1">${isUser ? 'You' : 'AI Assistant'}</div>
                <div class="p-3.5 rounded-2xl max-w-xl ${isUser ? 'bg-indigo-600 text-white' : 'bg-slate-900 text-slate-200 border border-slate-700'} markdown-body">
                    ${isUser ? content : marked.parse(content)}
                    ${sourcesHtml}
                </div>
            `;

            chatBox.appendChild(wrapper);
            chatBox.scrollTop = chatBox.scrollHeight;
        }

        fetchDocs();
    </script>
</body>
</html>
"""


@app.route('/')
def home():
    return render_template_string(HTML_TEMPLATE)


@app.route('/upload', methods=['POST'])
def upload_files():
    files = request.files.getlist('files')
    if not files:
        return jsonify({'error': 'No files provided'}), 400

    all_chunks = []
    for file_obj in files:
        filename = file_obj.filename
        if not filename:
            continue
        save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file_obj.save(save_path)

        chunks = extract_text_from_file(save_path, filename)
        if chunks:
            all_chunks.extend(chunks)
            uploaded_docs_registry.append({
                'filename': filename,
                'chunks': len(chunks),
                'path': save_path
            })

    if not all_chunks:
        return jsonify({'error': 'No readable text found in uploaded files'}), 400

    # Index into FAISS
    vectorstore = get_vectorstore()
    if vectorstore is None:
        vectorstore = FAISS.from_documents(all_chunks, embeddings_model)
    else:
        vectorstore.add_documents(all_chunks)

    vectorstore.save_local(app.config['FAISS_FOLDER'])

    return jsonify({
        'status': 'success',
        'total_chunks': len(all_chunks),
        'message': f'Successfully embedded {len(all_chunks)} chunks into FAISS vector database.'
    })


@app.route('/documents', methods=['GET'])
def list_documents():
    return jsonify({'documents': uploaded_docs_registry})


@app.route('/clear', methods=['POST'])
def clear_index():
    global uploaded_docs_registry
    uploaded_docs_registry = []
    if os.path.exists(app.config['FAISS_FOLDER']):
        shutil.rmtree(app.config['FAISS_FOLDER'])
    os.makedirs(app.config['FAISS_FOLDER'], exist_ok=True)
    return jsonify({'status': 'success', 'message': 'Index and documents cleared.'})


@app.route('/query', methods=['POST'])
def query_rag():
    data = request.get_json() or {}
    query_text = data.get('query', '').strip()
    if not query_text:
        return jsonify({'error': 'Query cannot be empty'}), 400

    vectorstore = get_vectorstore()
    if not vectorstore:
        return jsonify({
            'answer': 'Please upload at least one document to the Knowledge Base first before asking questions.',
            'sources': []
        })

    # Retrieve top 3 relevant chunks
    results_with_scores = vectorstore.similarity_search_with_score(query_text, k=3)

    sources = []
    context_blocks = []
    for idx, (doc, score) in enumerate(results_with_scores, 1):
        content = doc.page_content.strip()
        filename = doc.metadata.get('filename', 'Document')
        page = doc.metadata.get('page', 1)

        sources.append({
            'filename': filename,
            'page': page,
            'score': round(float(score), 4),
            'snippet': content[:250] + ('...' if len(content) > 250 else '')
        })

        context_blocks.append(f"[Excerpt {idx} | Source: {filename}, Page: {page}]\n{content}")

    context_str = "\n\n".join(context_blocks)

    system_prompt = (
        "You are an expert AI Document Assistant. "
        "Answer the question accurately using ONLY the provided document excerpts below. "
        "If the answer cannot be found in the context, state that the documents do not provide enough information."
    )

    user_prompt = f"DOCUMENT EXCERPTS:\n{context_str}\n\nUSER QUESTION:\n{query_text}"

    # Generate response via LLM
    answer = invoke_llm(user_prompt, system_prompt)

    # Fallback if no external LLM key is configured
    if not answer:
        answer = f"**Findings from your documents for \"{query_text}\":**\n\n"
        for idx, s in enumerate(sources, 1):
            answer += f"**{idx}. From {s['filename']} (Page {s['page']}):**\n> \"{s['snippet']}\"\n\n"
        answer += "---\n💡 *Tip: Add your `GEMINI_API_KEY` or run Ollama to enable conversational AI synthesis.*"

    return jsonify({
        'answer': answer,
        'sources': sources
    })


if __name__ == '__main__':
    print("\n" + "="*60)
    print("🚀 My RAG Document Assistant is running at: http://127.0.0.1:5000")
    print("="*60 + "\n")
    app.run(host='127.0.0.1', port=5000, debug=True)
